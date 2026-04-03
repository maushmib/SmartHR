from django.shortcuts import render, redirect
from django.contrib.auth import authenticate, login, logout
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.utils import timezone
from django.shortcuts import get_object_or_404
from .models import User, TeamLead, Employee, Project, Task
from .forms import CustomUserCreationForm, TeamLeadForm, EmployeeForm, ProjectForm, TaskAssignForm, TaskSubmitForm
def home(request):
    return render(request, 'core/home.html')
def register(request):
    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
           
            if user.role == "Guest":
                Candidate.objects.create(
                    user=user,
                    name=user.username,
                    email=user.email
                )
            return redirect('login')
    else:
        form = CustomUserCreationForm()
    return render(request, 'core/register.html', {'form': form})
def login_view(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return redirect('dashboard')
        else:
            messages.error(request, 'Invalid username or password')
    return render(request, 'core/login.html')
def logout_view(request):
    logout(request)
    return redirect('login')
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
@login_required
def dashboard(request):
    user = request.user
    if user.role == "HR":
     return hr_dashboard(request)
    elif user.role == "TL":
     return teamlead_dashboard(request)
    elif user.role == "EMP":
     return employee_dashboard(request)
    elif user.role == "Guest":
     return guest_dashboard(request)
    elif user.role == "Manager":
        return manager_dashboard(request)  

    else:
        messages.error(request, "Unknown role.")
        return redirect("login")
@login_required
def hr_dashboard(request):
    tasks = Task.objects.select_related('employee', 'team_lead', 'project')
    employee_id = request.GET.get('employee', 'all')
    team_lead_id = request.GET.get('team_lead', 'all')
    project_id = request.GET.get('project', 'all')
    submitted = request.GET.get('submitted', 'all')
    approved = request.GET.get('approved', 'all')
    date_start = request.GET.get('date_start')
    date_end = request.GET.get('date_end')
    if employee_id != 'all':
        tasks = tasks.filter(employee_id=employee_id)
    if team_lead_id != 'all':
        tasks = tasks.filter(team_lead_id=team_lead_id)
    if project_id != 'all':
        tasks = tasks.filter(project_id=project_id)
    if submitted == 'yes':
        tasks = tasks.filter(submitted=True)
    elif submitted == 'no':
        tasks = tasks.filter(submitted=False)
    if approved == 'yes':
        tasks = tasks.filter(approved=True)
    elif approved == 'no':
        tasks = tasks.filter(approved=False)
    if date_start:
        tasks = tasks.filter(date_assigned__gte=date_start)
    if date_end:
        tasks = tasks.filter(date_assigned__lte=date_end)
    context = {
        'tasks': tasks,
        'employees': Employee.objects.all(),
        'team_leads': TeamLead.objects.all(),
        'projects': Project.objects.all(),
        'filters': {
            'employee': employee_id,
            'team_lead': team_lead_id,
            'project': project_id,
            'submitted': submitted,
            'approved': approved,
            'date_start': date_start,
            'date_end': date_end
        }
    }
    return render(request, 'core/hr_dashboard.html', context)
@login_required
def employee_dashboard(request):
    try:
        employee = Employee.objects.get(user=request.user)
    except Employee.DoesNotExist:
        messages.error(request, "You are not an Employee.")
        return redirect('home')
    tasks = Task.objects.filter(employee=employee)
    profile_complete = all([
        employee.age,
        employee.gender,
        employee.education,
        employee.education_field,
        employee.job_role,
        employee.marital_status,
        employee.monthly_income
    ])
    if request.method == 'POST':
        form = TaskSubmitForm(request.POST)
        if form.is_valid():
            task_id = form.cleaned_data['task_id']
            try:
                task = Task.objects.get(id=task_id, employee=employee)
                task.submitted = True
                task.save()
                messages.success(request, "Task submitted for approval.")
            except Task.DoesNotExist:
                messages.error(request, "Task not found.")
            return redirect('dashboard')
    else:
        form = TaskSubmitForm()
    return render(request, 'core/employee_dashboard.html', {
        'tasks': tasks,
        'form': form,
        'employee': employee,
        'profile_complete': profile_complete  
    })
@login_required
def add_teamlead(request):
    form = TeamLeadForm(request.POST or None)
    if form.is_valid():
        form.save()
        return redirect('dashboard')
    return render(request, 'core/add_edit.html', {'form': form, 'title': 'Add Team Lead'})

@login_required
def add_employee(request):
    form = EmployeeForm(request.POST or None)
    if form.is_valid():
        form.save()  # This saves the employee with the selected team lead
        return redirect('dashboard')
    return render(request, 'core/add_edit.html', {'form': form, 'title': 'Add Employee'})
@login_required
def add_project(request):
    form = ProjectForm(request.POST or None)
    if form.is_valid():
        project = form.save(commit=False)
        # project.team_lead is already selected in form
        project.save()
        # Optional: link employees automatically (if you have a Project-Employee M2M)
        employees = Employee.objects.filter(team_lead=project.team_lead)
        project.employees.set(employees)  # if you have M2M
        return redirect('dashboard')
    return render(request, 'core/add_edit.html', {'form': form, 'title': 'Add Project'})

@login_required
def approve_task(request, task_id):
    task = get_object_or_404(Task, id=task_id)
    try:
        team_lead = TeamLead.objects.get(user=request.user)
    except TeamLead.DoesNotExist:
        messages.error(request, "You are not a Team Lead.")
        return redirect('dashboard')
    if task.team_lead != team_lead:
        messages.error(request, "You are not authorized to approve this task.")
        return redirect('dashboard')
    if request.method == 'POST':
        if task.submitted and not task.approved:
            task.approved = True
            task.date_submitted = timezone.now()  
            task.save()
            messages.success(request, "Task approved.")
        elif task.approved:
            messages.info(request, "Task already approved.")
        else:
            messages.error(request, "Task has not been submitted yet.")
    return redirect('dashboard')
from .models import Project, Task
from django import forms
from .models import Task, Employee
class TaskAssignForm(forms.ModelForm):
    project = forms.CharField(
        label='Project Name',
        max_length=100,
        widget=forms.TextInput(attrs={'class': 'form-control'})
    )
    class Meta:
        model = Task
        fields = ['employee', 'description']
    def __init__(self, *args, **kwargs):
        super(TaskAssignForm, self).__init__(*args, **kwargs)
        self.fields['employee'].widget.attrs.update({'class': 'form-control'})
        self.fields['description'].widget.attrs.update({'class': 'form-control'})
        if self.instance and self.instance.pk:
            self.fields['project'].initial = self.instance.project.name
@login_required
def delete_task(request, task_id):
    try:
        task = Task.objects.get(id=task_id)
        if task.team_lead.user == request.user:
            task.delete()
            messages.success(request, "Task deleted successfully.")
    except Task.DoesNotExist:
        messages.error(request, "Task does not exist.")
    return redirect('dashboard')
from django.utils import timezone
@login_required
def submit_task(request, task_id):
    try:
        employee = Employee.objects.get(user=request.user)
    except Employee.DoesNotExist:
        messages.error(request, "You are not a valid employee.")
        return redirect('dashboard')
    task = get_object_or_404(Task, id=task_id, employee=employee)
    if request.method == 'POST':
        if not task.submitted:
            task.submitted = True
            task.save()
            messages.success(request, "Task submitted successfully. Waiting for Team Lead approval.")
        else:
            messages.info(request, "Task already submitted.")
    return redirect('employee_dashboard')
from django.shortcuts import render, redirect
from .forms import TaskAssignForm
from .models import Task
@login_required
def teamlead_dashboard(request):
    try:
        teamlead = TeamLead.objects.get(user=request.user)
    except TeamLead.DoesNotExist:
        messages.error(request, "You are not a Team Lead.")
        return redirect('home')
    tasks = Task.objects.filter(team_lead=teamlead)
    employees = Employee.objects.filter(team_lead=teamlead)
    projects = Project.objects.filter(team_lead=teamlead)
    
    if request.method == 'POST':
        form = TaskAssignForm(request.POST)
        form.fields['employee'].queryset = employees
        form.fields['project'].queryset = projects
        if form.is_valid():
            task = form.save(commit=False)
            task.team_lead = teamlead  
            task.save()
            messages.success(request, "Task assigned successfully.")
            return redirect('dashboard')
    else:
        form = TaskAssignForm()
        form.fields['employee'].queryset = employees
        form.fields['project'].queryset = projects
        
    return render(request, 'core/teamlead_dashboard.html', {
        'form': form,
        'tasks': tasks,
        'employees': employees,
        'projects': projects,
    })
from .models import TeamLead
from django.shortcuts import render, redirect
from .models import Project, Employee, TeamLead
from .forms import TaskAssignForm
from django.http import HttpResponse
from .models import Employee, Task
from django.shortcuts import render, redirect
from .models import Project, Employee, TeamLead, Task
from .forms import TaskAssignForm

def assign_task(request):
    try:
        team_lead = TeamLead.objects.get(user=request.user)
    except TeamLead.DoesNotExist:
        return HttpResponse("You are not a valid team lead.", status=403)

    # Only employees under this TL
    employees = Employee.objects.filter(team_lead=team_lead)
    
    # Only projects assigned to this TL
    projects = Project.objects.filter(team_lead=team_lead)

    if request.method == 'POST':
        form = TaskAssignForm(request.POST)
        if form.is_valid():
            task = form.save(commit=False)
            task.team_lead = team_lead
            task.save()
            return redirect('teamlead_dashboard')
    else:
        form = TaskAssignForm()

    return render(request, 'core/assign_task.html', {
        'form': form,
        'employees': employees,
        'projects': projects
    })


@login_required
def edit_task(request, task_id):
    task = get_object_or_404(Task, id=task_id)
    try:
        teamlead = TeamLead.objects.get(user=request.user)
    except TeamLead.DoesNotExist:
        messages.error(request, "You are not a Team Lead.")
        return redirect('home')
    employees = Employee.objects.filter(team_lead=teamlead)
    projects = Project.objects.filter(team_lead=teamlead)
    if request.method == 'POST':
        form = TaskAssignForm(request.POST, instance=task)
        form.fields['employee'].queryset = employees
        form.fields['project'].queryset = projects
        if form.is_valid():
            edited_task = form.save(commit=False)
            edited_task.team_lead = teamlead 
            edited_task.save()
            messages.success(request, "Task updated successfully.")
            return redirect('dashboard')
    else:
        form = TaskAssignForm(instance=task)
        form.fields['employee'].queryset = employees
        form.fields['project'].queryset = projects
    return render(request, 'core/add_edit.html', {'form': form, 'title': 'Edit Task'})
from django.contrib.auth import logout
from django.shortcuts import redirect
def logout_view(request):
    logout(request)
    return redirect('login')  
from django.shortcuts import render, redirect
from .forms import ProjectForm
from django.shortcuts import render, redirect
from .forms import ProjectForm
def create_project(request):
    if request.method == 'POST':
        form = ProjectForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('teamlead_dashboard')  
    else:
        form = ProjectForm()
    return render(request, 'core/create_project.html', {'form': form})
from django.contrib.auth.decorators import login_required
from .forms import PerformanceRatingForm
from .models import PerformanceRating
from django.shortcuts import render, redirect
@login_required
def give_rating(request):
    if request.user.role != 'TL':
        return redirect('unauthorized')
    if request.method == 'POST':
        form = PerformanceRatingForm(request.POST)
        if form.is_valid():
            rating = form.save(commit=False)
            rating.rated_by = request.user
            rating.save()
            return redirect('rating_success')
    else:
        form = PerformanceRatingForm()
    return render(request, 'core/give_rating.html', {'form': form})
from .models import PerformanceRating
@login_required
def view_my_ratings(request):
    if request.user.role != 'EMP':
        return redirect('unauthorized')
    try:
        employee = Employee.objects.get(user=request.user)
    except Employee.DoesNotExist:
        messages.error(request, "You are not a valid employee.")
        return redirect('dashboard')
    ratings = PerformanceRating.objects.filter(employee=employee)
    return render(request, 'core/view_my_ratings.html', {'ratings': ratings})
@login_required
def view_all_ratings(request):
    if request.user.role != 'HR':
        return redirect('unauthorized')
    ratings = PerformanceRating.objects.all().order_by('-date_rated')
    return render(request, 'core/all_ratings.html', {'ratings': ratings})
from django.shortcuts import render
def rating_success(request):
    return render(request, 'core/rating_success.html')
import os, docx, PyPDF2
from django.conf import settings
from .models import Employee
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import os, docx, PyPDF2
from django.conf import settings
from .models import Employee
def _extract_pdf(path):
    text = ""
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() or ""
    return text
def _extract_docx(path):
    doc = docx.Document(path)
    return "\n".join(p.text for p in doc.paragraphs)
import re, nltk
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
stop_words  = set(stopwords.words("english"))
lemmatizer  = WordNetLemmatizer()
tokenizer   = nltk.word_tokenize
def keywords(text):
    words = tokenizer(text.lower())
    return {lemmatizer.lemmatize(w) for w in words if w.isalnum() and w not in stop_words}
def overlap_score(jd_text, resume_text):
    jd_keywords      = keywords(jd_text)
    resume_keywords  = keywords(resume_text)
    if not jd_keywords:
        return 0.0
    common = jd_keywords & resume_keywords
    return round(len(common) / len(jd_keywords) * 100, 2)
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
model = SentenceTransformer("all-MiniLM-L6-v2")   
def semantic_score(jd_text, resume_text):
    emb_jd      = model.encode([jd_text])[0]
    emb_resume  = model.encode([resume_text])[0]
    sim         = cosine_similarity([emb_jd], [emb_resume])[0][0]
    return round(sim * 100, 2)
@login_required
def resume_match_view(request):
    if request.user.role != "HR":
        return redirect("unauthorized")
    employees = Employee.objects.all()
    match_score = None
    chosen_emp = None
    if request.method == "POST":
        emp_id = request.POST.get("employee_id")
        jd_text = request.POST.get("jd", "")
        print("✅ Received employee_id from form:", emp_id)
        if not emp_id:
            messages.error(request, "No employee selected.")
            return render(request, "core/resume_match.html", {"employees": employees})
        chosen_emp = get_object_or_404(Employee, id=emp_id)
        resume_path = os.path.join(settings.MEDIA_ROOT, chosen_emp.resume_file.name)
        ext = chosen_emp.resume_file.name.split(".")[-1].lower()
        if ext in ("doc", "docx"):
            resume_text = _extract_docx(resume_path)
        elif ext == "pdf":
            resume_text = _extract_pdf(resume_path)
        else:
            messages.error(request, "Unsupported resume file type.")
            return render(request, "core/resume_match.html", {"employees": employees})
        match_score = overlap_score(jd_text, resume_text)
    return render(
        request,
        "core/resume_match.html",
        {
            "employees": employees,
            "chosen_emp": chosen_emp,
            "match_score": match_score,
        },
    )
@login_required
def upload_resume(request):
    if request.user.role != 'EMP':
        return redirect('unauthorized')
    employee = Employee.objects.get(user=request.user)
    if request.method == 'POST':
        uploaded = request.FILES.get('resume_file')
        if uploaded:
            employee.resume_file = uploaded
            employee.save()
            messages.success(request, "Resume uploaded successfully.")
            return redirect('employee_dashboard')
    return render(request, 'core/upload_resume.html', {'employee': employee})
import json
import requests
import subprocess
from django.shortcuts import render
from .models import Employee
import sqlite3
from django.conf import settings
import sqlite3
from .forms import HRFilterForm
from .models import Task, PerformanceRating
from django.shortcuts import render
from .models import Employee, TeamLead, Project, Task, PerformanceRating
from django.shortcuts import render
from .models import Employee, TeamLead, Project, Task
def filter_page(request):
    employees = Employee.objects.all()
    team_leads = TeamLead.objects.all()
    projects = Project.objects.all()
    filtered_tasks = None
    if request.method == 'POST':
        employee_id = request.POST.get('employee_id')
        team_lead_id = request.POST.get('team_lead_id')
        project_id = request.POST.get('project_id')
        submitted = request.POST.get('submitted')
        approved = request.POST.get('approved')
        filters = {}
        if employee_id:
            filters['employee_id'] = employee_id
        if team_lead_id:
            filters['employee__team_lead_id'] = team_lead_id
        if project_id:
            filters['project_id'] = project_id
        if submitted != "":
            filters['submitted'] = submitted
        if approved != "":
            filters['approved'] = approved
        filtered_tasks = Task.objects.filter(**filters)
    return render(request, 'core/filter_page.html', {
        'employees': employees,
        'team_leads': team_leads,
        'projects': projects,
        'filtered_tasks': filtered_tasks
    })
def filter_results(request):
    if request.method == 'POST':
        employee_id = request.POST.get('employee_id')
        team_lead_id = request.POST.get('team_lead_id')
        project_id = request.POST.get('project_id')
        submitted = request.POST.get('submitted')
        approved = request.POST.get('approved')
        filters = {}
        if employee_id:
            filters['employee_id'] = employee_id
        if team_lead_id:
            filters['employee__team_lead_id'] = team_lead_id
        if project_id:
            filters['project_id'] = project_id
        if submitted != "":
            filters['submitted'] = submitted
        if approved != "":
            filters['approved'] = approved
        filtered_tasks = Task.objects.filter(**filters)
        return render(request, 'core/filter_results.html', {
            'tasks': filtered_tasks
        })
    return render(request, 'core/filter_results.html', {'tasks': []})
def predict_attrition(request, employee_id):
    employee = get_object_or_404(Employee, id=employee_id)
    required_fields = [
        employee.age, employee.gender, employee.education,
        employee.education_field, employee.job_role,
        employee.marital_status, employee.monthly_income
    ]
    if not all(required_fields):
        return render(request, 'error.html', {'message': 'Please complete your profile to get a prediction.'})
    input_data = pd.DataFrame([{
        'Age': employee.age,
        'Gender': employee.gender,
        'Education': employee.education,
        'EducationField': employee.education_field,
        'JobRole': employee.job_role,
        'MaritalStatus': employee.marital_status,
        'MonthlyIncome': employee.monthly_income,
    }])
    model = joblib.load('core/ml/attrition_model.pkl')
    prediction = model.predict(input_data)[0]
    return render(request, 'attrition_result.html', {
        'employee': employee,
        'prediction': 'Likely to Leave' if prediction == 1 else 'Likely to Stay'
    })
from django.contrib.auth.decorators import login_required
from django.shortcuts import render
from .models import Employee
import pandas as pd
import joblib
import os

# -----------------------------
# Load ML model, scaler, features
# -----------------------------
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MODEL_PATH = r'D:\internship\hr_management\core\ml\attrition_model.pkl'
SCALER_PATH = r'D:\internship\hr_management\core\ml\scaler.pkl'
FEATURES_PATH = r'D:\internship\hr_management\core\ml\model_features.pkl'

model = joblib.load(MODEL_PATH)
scaler = joblib.load(SCALER_PATH)
model_features = joblib.load(FEATURES_PATH)

@login_required
def attrition_prediction_view(request):
    employees = Employee.objects.all()
    predictions = []

    # Prepare data for ML model
    data = []
    for emp in employees:
        data.append({
            "Employee Name": emp.name,
            "Age": emp.age,
            "Gender": emp.gender,
            "Education": emp.education,
            "EducationField": emp.education_field,
            "JobRole": emp.job_role,
            "MonthlyIncome": emp.monthly_income
        })

    df = pd.DataFrame(data)
    
    if not df.empty:
        # Feature engineering: Salary_vs_Role_Avg
        role_avg_salary = df.groupby('JobRole')['MonthlyIncome'].transform('mean')
        df['Salary_vs_Role_Avg'] = df['MonthlyIncome'] / role_avg_salary

        # One-hot encode categorical columns
        df_encoded = pd.get_dummies(df, columns=['Gender','EducationField','JobRole'], drop_first=True)

        # Add missing columns (0)
        for col in model_features:
            if col not in df_encoded.columns:
                df_encoded[col] = 0

        # Ensure column order
        df_encoded = df_encoded[model_features]

        # Scale numeric features
        numeric_cols = ['Age', 'MonthlyIncome', 'Education', 'Salary_vs_Role_Avg']
        df_encoded[numeric_cols] = scaler.transform(df_encoded[numeric_cols])

        # Predict
        probs = model.predict_proba(df_encoded)[:,1]
        threshold = 0.35
        y_pred = (probs > threshold).astype(int)

        for i, emp in enumerate(employees):
            predictions.append({
                'name': emp.name,
                'department': emp.department,
                'job_role': emp.job_role,
                'monthly_income': emp.monthly_income,
                'prediction': 'Yes' if y_pred[i]==1 else 'No'
            })

    return render(request, 'core/attrition_prediction.html', {'predictions': predictions})

@login_required
def edit_employee_profile(request):
    try:
        employee = Employee.objects.get(user=request.user)
    except Employee.DoesNotExist:
        messages.error(request, "Employee profile not found.")
        return redirect('dashboard')
    if request.method == 'POST':
        form = EmployeeForm(request.POST, instance=employee)
        if form.is_valid():
            form.save()
            messages.success(request, "Profile updated successfully.")
            return redirect('dashboard')
    else:
        form = EmployeeForm(instance=employee)
    return render(request, 'core/edit_employee_profile.html', {'form': form})
from django.shortcuts import render
from .models import TeamLead, Employee, Task, Project
from joblib import load
import os
from django.shortcuts import render
from .models import TeamLead, Employee, Task, Project
from joblib import load
import os
from django.shortcuts import render
from .models import TeamLead, Project, Employee, Task
from joblib import load
import os
import os
import joblib
from django.conf import settings
from .models import Employee, Project
from django.shortcuts import render
import os
import joblib
from django.shortcuts import render
from .models import Employee
import os
import joblib
from django.shortcuts import render
from .models import Employee, Project, TeamLead, Task
from django.shortcuts import render
from .models import Employee, Task
import joblib
import os
import pandas as pd
import joblib
from django.shortcuts import render
from .models import Employee, Task
from django.shortcuts import render
from .models import Project, Employee, TeamLead
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

def smart_assign(request):
    try:
        team_lead = TeamLead.objects.get(user=request.user)
        employees = Employee.objects.filter(team_lead=team_lead)
    except TeamLead.DoesNotExist:
        employees = []

    projects = Project.objects.filter(team_lead=team_lead)
    distinct_projects = projects.values_list('name', flat=True).distinct()
    recommendations = []
    selected_project = None

    if request.method == 'POST':
        selected_project_name = request.POST.get('project_name')
        if selected_project_name:
            selected_project = Project.objects.get(name=selected_project_name, team_lead=team_lead)
            if selected_project.required_skills:
                project_skills_text = selected_project.required_skills.lower()

                for emp in employees:
                    if emp.skills:
                        emp_skills_text = emp.skills.lower()

                        # Convert skills to TF-IDF vectors
                        vectorizer = TfidfVectorizer()
                        vectors = vectorizer.fit_transform([project_skills_text, emp_skills_text])

                        # Compute cosine similarity
                        sim_score = cosine_similarity(vectors[0:1], vectors[1:2])[0][0]
                        score = int(sim_score * 100)  # Convert to percentage

                        if score > 0:
                            recommendations.append({
                                'employee': emp,
                                'score': score
                            })

    # Sort recommendations by score descending
    recommendations.sort(key=lambda x: x['score'], reverse=True)

    return render(request, 'core/smart_assign.html', {
        'distinct_projects': distinct_projects,
        'recommendations': recommendations,
        'selected_project': selected_project
    })



from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
from .models import Employee, Project
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
@login_required
def skill_gap_analyzer(request, employee_id, project_id):
    employee = get_object_or_404(Employee, id=employee_id)
    project = get_object_or_404(Project, id=project_id)
    employee_skills = set(s.strip().lower() for s in (employee.skills or "").split(',') if s.strip())
    project_skills = set(s.strip().lower() for s in (project.required_skills or "").split(',') if s.strip())
    missing_skills = project_skills - employee_skills
    training_courses = {
        "python": {"name": "Python Programming Basics", "provider": "Coursera", "url": "https://www.coursera.org/specializations/python"},
        "django": {"name": "Django for Beginners", "provider": "Udemy", "url": "https://www.udemy.com/course/django-for-beginners/"},
        "cybersecurity": {"name": "Cybersecurity Basics", "provider": "edX", "url": "https://www.edx.org/course/cybersecurity-basics"},
        "sql": {"name": "SQL for Data Analysis", "provider": "DataCamp", "url": "https://www.datacamp.com/courses/sql-for-data-analysis"},
        "react": {"name": "React - The Complete Guide", "provider": "Udemy", "url": "https://www.udemy.com/course/react-the-complete-guide/"},
        "java": {"name": "Java Programming Masterclass", "provider": "Udemy", "url": "https://www.udemy.com/course/java-the-complete-java-developer-course/"},
        "data science": {"name": "Data Science Specialization", "provider": "Coursera", "url": "https://www.coursera.org/specializations/jhu-data-science"},
        "machine learning": {"name": "Machine Learning by Andrew Ng", "provider": "Coursera", "url": "https://www.coursera.org/learn/machine-learning"},
        "cloud computing": {"name": "Cloud Computing Basics", "provider": "edX", "url": "https://www.edx.org/course/cloud-computing-basics"},
        "scikitlearn": {"name": "Scikitlearn- Machine Learning with Python", "provider": "Coursera", "url": "https://www.coursera.org/learn/python-machine-learning"},
        "devops": {"name": "DevOps Essentials", "provider": "Udemy", "url": "https://www.udemy.com/course/devops-foundations/"}
    }
    recommended_courses = []
    if missing_skills:
        for skill in missing_skills:
            corpus = [skill] + [course["name"] for course in training_courses.values()]
            vectorizer = TfidfVectorizer().fit_transform(corpus)
            vectors = vectorizer.toarray()
            similarity_scores = cosine_similarity([vectors[0]], vectors[1:])[0]
            best_idx = similarity_scores.argmax()
            best_course_key = list(training_courses.keys())[best_idx]
            best_score = similarity_scores[best_idx]
            if best_score > 0.2:
                best_course = training_courses[best_course_key].copy()
                best_course["score"] = float(best_score)
                recommended_courses.append(best_course)
    context = {
        'employee': employee,
        'project': project,
        'missing_skills': missing_skills,
        'recommended_courses': recommended_courses,
    }
    return render(request, 'core/skill_gap_results.html', context)
import os
import joblib
BASE_DIR = os.path.dirname(os.path.abspath(__file__)) 
model_path = os.path.join(BASE_DIR, 'ml', 'resume_model.pkl')
vectorizer_path = os.path.join(BASE_DIR, 'ml', 'resume_vectorizer.pkl')
resume_model = joblib.load(model_path)
resume_vectorizer = joblib.load(vectorizer_path)
def clean_text(text):
    text = re.sub(r'http\S+', '', text)
    text = re.sub(r'[^a-zA-Z]', ' ', text)
    return text.lower()
def predict_resume_authenticity(resume_text):
    cleaned = clean_text(resume_text)
    X_vec = resume_vectorizer.transform([cleaned])
    pred = resume_model.predict(X_vec)[0]
    return "Real / Genuine" if pred == 1 else "Fake / Irrelevant"
def guest_resume_submit(request):
    if request.method == "POST":
        name = request.POST.get("name")
        email = request.POST.get("email")
        resume_file = request.FILES.get("resume_file")   
        if not name or not email or not resume_file:
            return render(request, "core/guest_resume_submit.html", {"error": "All fields are required"})
        candidate = Candidate.objects.create(
            user=request.user,
            name=name,
            email=email,
            resume_file=resume_file,
            status="Pending",
        )
        resume_path = os.path.join(settings.MEDIA_ROOT, candidate.resume_file.name)
        ext = candidate.resume_file.name.split(".")[-1].lower()
        resume_text = ""
        try:
            if ext == "pdf":
                import PyPDF2
                with open(resume_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    resume_text = "".join([p.extract_text() or "" for p in reader.pages])
            elif ext == "docx":
                import docx
                doc = docx.Document(resume_path)
                resume_text = "\n".join(p.text for p in doc.paragraphs)
        except:
            pass
        if resume_text:
            candidate.status = predict_resume_authenticity(resume_text)
            candidate.save()
        return render(request, "core/guest_resume_submit.html", {"success": "Resume submitted successfully!"})
    return render(request, "core/guest_resume_submit.html")
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect
from django.contrib import messages
from .models import Candidate
import os, re
from django.conf import settings
import joblib
model_path = os.path.join(settings.BASE_DIR, "core", "ml", "resume_model.pkl")
vectorizer_path = os.path.join(settings.BASE_DIR, "core", "ml", "resume_vectorizer.pkl")
resume_model = joblib.load(model_path)
resume_vectorizer = joblib.load(vectorizer_path)
def clean_text(text):
    text = re.sub(r'http\S+', '', text)
    text = re.sub(r'[^a-zA-Z]', ' ', text)
    return text.lower()
def predict_resume_authenticity(resume_text):
    cleaned = clean_text(resume_text)
    X_vec = resume_vectorizer.transform([cleaned])
    pred = resume_model.predict(X_vec)[0]
    return "Approved" if pred == 1 else "Rejected"
@login_required
def hr_candidates(request):
    if request.user.role != "HR":
        return redirect("dashboard")
    candidates = Candidate.objects.all()
    for c in candidates:
        if c.resume_file:
            resume_path = os.path.join(settings.MEDIA_ROOT, c.resume_file.name)
            try:  
                ext = c.resume_file.name.split(".")[-1].lower()
                if ext in ("pdf", "docx"):
                    if ext == "pdf":
                        import PyPDF2
                        with open(resume_path, "rb") as f:
                            reader = PyPDF2.PdfReader(f)
                            resume_text = "".join([p.extract_text() or "" for p in reader.pages])
                    else:  
                        import docx
                        doc = docx.Document(resume_path)
                        resume_text = "\n".join(p.text for p in doc.paragraphs)
                else:
                    resume_text = ""
                c.status = predict_resume_authenticity(resume_text)
                c.save()
            except:
                c.status = "Error reading file"
        else:
            c.status = "No resume"
    return render(request, "core/hr_candidates.html", {"candidates": candidates})
@login_required
def guest_dashboard(request):
    try:
        candidate = Candidate.objects.get(user=request.user)
    except Candidate.DoesNotExist:
        messages.error(request, "Candidate not found.")
        return redirect("home")
    if request.method == "POST" and request.FILES.get("resume_file"):
        resume_file = request.FILES["resume_file"]
        candidate.resume_file = resume_file
        candidate.status = "Pending"
        candidate.save()
        resume_path = os.path.join(settings.MEDIA_ROOT, candidate.resume_file.name)
        ext = candidate.resume_file.name.split(".")[-1].lower()
        resume_text = ""
        try:
            if ext == "pdf":
                import PyPDF2
                with open(resume_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    resume_text = "".join([p.extract_text() or "" for p in reader.pages])
            elif ext == "docx":
                import docx
                doc = docx.Document(resume_path)
                resume_text = "\n".join(p.text for p in doc.paragraphs)
        except:
            resume_text = ""
        if resume_text:
            candidate.status = predict_resume_authenticity(resume_text)
            candidate.save()
        messages.success(request, "Resume uploaded & processed successfully!")
    return render(request, "core/guest_dashboard.html", {"candidate": candidate})
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator
from django.views import View
import json
from .models import Employee, TeamLead, Project, Task, Candidate, PerformanceRating
from ollama import Client
ollama_client = Client()
MODEL_MAP = {
    "Employee": Employee,
    "TeamLead": TeamLead,
    "Project": Project,
    "Task": Task,
    "Candidate": Candidate,
    "PerformanceRating": PerformanceRating
}
SYNONYM_MAP = {
    "authenticity": "status",
    "resume": "status",
    "candidates names": "name",
    "employee skills": "skills",
    "task assigned": "description",
    "teamlead": "team_lead__name",     
    "team lead": "team_lead__name"     
}
@method_decorator(csrf_exempt, name='dispatch')
class HRChatbot(View):
    def post(self, request, *args, **kwargs):
        try:
            data = json.loads(request.body)
            question = data.get("question", "").strip()
            if not question:
                return JsonResponse({"answer": "Please ask a question."})
            model_info = """
Employee: name, email, skills, team_lead (FK), age, gender, education, education_field, job_role, marital_status, monthly_income, department
TeamLead: name, email, user (FK)
Project: name, description, team_lead (FK), domain, required_skills, deadline, criticality
Task: project (FK), employee (FK), team_lead (FK), description, date_assigned, submitted, approved, date_submitted
Candidate: user (FK), name, email, resume_file, status, prediction, notified
PerformanceRating: employee (FK), rated_by (FK), rating, comments, date_rated
"""
            prompt = f"""
You are an HR assistant. You have access to these Django models and fields:
{model_info}
User question: "{question}"
Instructions:
- Respond ONLY in valid JSON.
- JSON keys: model, filter, fields
- Use Django ORM syntax for filters.
- For text fields (skills, job_role, department, status), use __icontains for partial match.
- For boolean fields, use true/false.
- Example: {{"model": "Employee", "filter": {{"skills__icontains": "Python"}}, "fields": ["name","email"]}}
- Respond strictly in JSON, no explanations.
"""
            response = ollama_client.chat(
                model="mistral",
                messages=[{"role": "user", "content": prompt}]
            )
            try:
                instructions_str = response['message']['content']
                ai_instructions = json.loads(instructions_str)
                model_name = ai_instructions.get("model")
                filters = ai_instructions.get("filter", {})
                fields = ai_instructions.get("fields", [])
            except Exception:
                model_name = None
                filters = {}
                fields = []
            for key in list(filters.keys()):
                for syn, actual in SYNONYM_MAP.items():
                    if syn in key:
                        filters[actual] = filters.pop(key)
            if not model_name:
                q_lower = question.lower()
                if "task" in q_lower:
                    model_name = "Task"
                    fields = ["employee__name","project__name","team_lead__name","description","date_assigned","date_submitted","approved"]
                    if "submitted" in q_lower:
                        filters["submitted"] = True
                    if "not submitted" in q_lower:
                        filters["submitted"] = False
                    if "approved" in q_lower:
                        filters["approved"] = True
                    if "not approved" in q_lower:
                        filters["approved"] = False
                elif "candidate" in q_lower or "resume" in q_lower:
                    model_name = "Candidate"
                    fields = ["name","email","status"]
                    if "rejected" in q_lower:
                        filters["status__icontains"] = "rejected"
                    elif "approved" in q_lower:
                        filters["status__icontains"] = "approved"
                elif "project" in q_lower:
                    model_name = "Project"
                    fields = ["name","description","team_lead__name","domain","deadline","criticality"]
                elif "employee" in q_lower:
                      model_name = "Employee"
                      fields = ["name","email","skills","team_lead__name"]
                      if "under teamlead" in q_lower:
                           filters["team_lead__name__icontains"] = q_lower.split("under teamlead")[-1].strip()
                      elif "under" in q_lower:
                           filters["team_lead__name__icontains"] = q_lower.split("under")[-1].strip()

            if not model_name or model_name not in MODEL_MAP:
                return JsonResponse({"answer": "Could not determine a valid model from your query."})
            queryset = MODEL_MAP[model_name].objects.filter(**filters).values(*fields)
            result_list = list(queryset)
            if not result_list:
                return JsonResponse({"answer": "No records found matching your query."})
            formatted_result = []
            for record in result_list:
                line = ", ".join([f"{k}: {v}" for k, v in record.items()])
                formatted_result.append(line)
            answer_text = "\n".join(formatted_result)
            return JsonResponse({"answer": answer_text})
        except Exception as e:
            return JsonResponse({"answer": f"Error: {str(e)}"})
        from django.shortcuts import render, redirect
from .models import Project, TeamLead
from .forms import ProjectForm

def create_project(request):
    if request.method == 'POST':
        form = ProjectForm(request.POST)
        if form.is_valid():
            form.save()  # project automatically has a team_lead field
            return redirect('manager_dashboard')
    else:
        form = ProjectForm()
    return render(request, 'core/create_project.html', {'form': form})

from django.shortcuts import render
from .models import Project, TeamLead, Task
from django.contrib.auth.decorators import login_required
from django.shortcuts import render
from .models import Project, TeamLead, Task
from django.shortcuts import render, redirect
from .models import TeamLead, Employee, Project
from .forms import AllocateEmployeeForm, ProjectForm
from django.http import HttpResponse
from django.shortcuts import render, redirect
from .models import TeamLead, Employee
from .forms import AllocateEmployeeForm  # simple form with employee and TL
# views.py
from django.shortcuts import redirect, render
from .models import TeamLead, Employee
from .forms import AllocateEmployeeForm
def allocate_employees(request):
    if request.method == 'POST':
        form = AllocateEmployeeForm(request.POST)
        if form.is_valid():
            # get the selected TeamLead from form
            tl = form.cleaned_data['team_lead']

            # assign employees to this TeamLead
            selected_emps = form.cleaned_data['employees']
            for emp in selected_emps:
                emp.team_lead = tl
                emp.save()

            return redirect('manager_dashboard')
    else:
        form = AllocateEmployeeForm()
    
    return render(request, "core/allocate_employees.html", {'form': form})


from django.shortcuts import render, redirect
from .models import Project, TeamLead
from django import forms

# Form for Manager to create project
class ManagerProjectForm(forms.ModelForm):
    class Meta:
        model = Project
        fields = ['name', 'description', 'domain', 'required_skills', 'deadline', 'criticality', 'team_lead']
from django.shortcuts import render, redirect
from .forms import ProjectForm

def manager_create_project(request):
    if request.method == "POST":
        form = ProjectForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('manager_dashboard')  # Redirect to dashboard after creation
    else:
        form = ProjectForm()
    return render(request, "core/create_project.html", {"form": form})

from django.shortcuts import render
from .models import Project, TeamLead

def choose_project(request):
    try:
        team_lead = TeamLead.objects.get(user=request.user)
        projects = Project.objects.filter(team_lead=team_lead)
    except TeamLead.DoesNotExist:
        projects = []

    return render(request, 'core/choose_project.html', {'projects': projects})
# views.py
from django.shortcuts import render
from .models import Project, Task, TeamLead

def manager_dashboard(request):
    tasks = Task.objects.select_related('project', 'employee', 'team_lead').all()
    team_leads = TeamLead.objects.prefetch_related('employee_set')  # fetch employees for TLs
    return render(request, "core/manager_dashboard.html", {
        "tasks": tasks,
        "team_leads": team_leads
    })
